import os
import json
import threading
import tkinter as tk
from tkinter import filedialog, messagebox
import customtkinter as ctk
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Config paths
APP_DIR = os.path.join(os.getenv('APPDATA', os.path.expanduser('~')), 'CER_Automator')
CONFIG_FILE = os.path.join(APP_DIR, 'config.json')

def load_settings():
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            pass
    return {}

def save_settings(settings: dict):
    os.makedirs(APP_DIR, exist_ok=True)
    try:
        with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(settings, f, indent=4)
    except Exception:
        pass

# Global Keywords for both parsing and replacing
SECTION_KEYWORDS = {
    "Contexte": ["contexte"],
    "Problématique": ["problématique", "problematique", "problematiques", "problématiques"],
    "Mots-Clefs": ["mots-clefs", "mots clefs", "mots-clés", "mots clés", "mot clef", "mot clé"],
    "Contrainte": ["contrainte", "contraintes"],
    "Hypothèses": ["hypothèses", "hypotheses", "hypothèse", "hypothese"],
    "Plan d’action": ["plan d’action", "plan d'action", "plan action", "étapes", "etapes"]
}

EXCLUDED_KEYWORDS = ["philosophie", "généralisation", "generalisation", "livrable", "livrables"]

def is_header(text, keywords, strict=False):
    """
    Checks if a text is a section header.
    If strict is True, it requires a colon or an exact match.
    """
    clean = text.lower().replace(" ", " ").strip()
    has_colon = ":" in clean
    no_colon = clean.replace(":", "").strip()
    
    for kw in keywords:
        # Match exact (e.g., "Contexte")
        if no_colon == kw:
            return True, has_colon
        # Match startswith only if followed by colon (e.g., "Contexte : ...")
        # OR if it's very short (e.g., "Contexte d'auth" -> but we want to avoid this)
        if clean.startswith(kw) and has_colon:
            # Check if the part before the colon is essentially the keyword
            before_colon = clean.split(":")[0].strip()
            if before_colon == kw:
                return True, True
    return False, False

def replace_content_by_titles(doc, replacements):
    """
    Finds section titles (e.g., 'Contexte :') and replaces the content of the 
    paragraphs immediately following them.
    """
    paragraphs = doc.paragraphs
    # Map to store the best paragraph index for each section
    # { section_name: (index, has_colon) }
    best_anchors = {}

    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text or len(text) > 60: continue

        for section_name, keywords in SECTION_KEYWORDS.items():
            found, with_colon = is_header(text, keywords)
            if found:
                # If we have no anchor yet, or if this one is better (has colon)
                if section_name not in best_anchors or (with_colon and not best_anchors[section_name][1]):
                    best_anchors[section_name] = (i, with_colon)

    # Apply replacements at best anchors
    for section_name, (anchor_idx, _) in best_anchors.items():
        if section_name in replacements and replacements[section_name]:
            # Find the next non-empty paragraph to replace
            j = anchor_idx + 1
            while j < len(paragraphs):
                if paragraphs[j].text.strip():
                    paragraphs[j].text = replacements[section_name]
                    paragraphs[j].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    break
                j += 1

class CERAutomatorApp(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent, fg_color="transparent")
        self.controller = controller

        settings = load_settings()
        self.template_path = ctk.StringVar(value=settings.get("template_path", ""))
        self.default_filename = ctk.StringVar(value=settings.get("default_filename", "CER_Genere"))
        
        # Fields for the 6 sections
        self.sections = {s: tk.StringVar() for s in SECTION_KEYWORDS.keys()}
        
        # Text widgets references
        self.text_widgets = {}

        self._build_ui()

    def _build_ui(self):
        # Header
        header_frame = ctk.CTkFrame(self, fg_color="transparent")
        header_frame.pack(fill="x", padx=20, pady=(15, 5))
        
        self.back_btn = ctk.CTkButton(
            header_frame, text="🏠 Menu", width=80, 
            command=lambda: self.controller.show_menu()
        )
        self.back_btn.pack(side="left")

        self.title_label = ctk.CTkLabel(
            header_frame, text="CERviable", 
            font=ctk.CTkFont(size=22, weight="bold")
        )
        self.title_label.pack(side="left", expand=True, padx=(0, 20))

        # Config Selection
        config_frame = ctk.CTkFrame(self)
        config_frame.pack(fill="x", padx=20, pady=10)
        
        template_row = ctk.CTkFrame(config_frame, fg_color="transparent")
        template_row.pack(fill="x", padx=10, pady=5)
        ctk.CTkLabel(template_row, text="Template :", font=ctk.CTkFont(weight="bold"), width=100).pack(side="left")
        self.template_entry = ctk.CTkEntry(template_row, textvariable=self.template_path, state="disabled")
        self.template_entry.pack(side="left", padx=10, fill="x", expand=True)
        ctk.CTkButton(template_row, text="Choisir", command=self._select_template, width=80).pack(side="left")

        filename_row = ctk.CTkFrame(config_frame, fg_color="transparent")
        filename_row.pack(fill="x", padx=10, pady=5)
        ctk.CTkLabel(filename_row, text="Nom Sortie :", font=ctk.CTkFont(weight="bold"), width=100).pack(side="left")
        self.filename_entry = ctk.CTkEntry(filename_row, textvariable=self.default_filename, placeholder_text="Ex: CER_Genere")
        self.filename_entry.pack(side="left", padx=10, fill="x", expand=True)
        self.default_filename.trace_add("write", self._save_filename_setting)

        # Prosit Selection
        prosit_frame = ctk.CTkFrame(self, fg_color="transparent")
        prosit_frame.pack(fill="x", padx=20, pady=5)
        
        self.load_prosit_btn = ctk.CTkButton(
            prosit_frame, text="📂 Charger un Prosit Aller (.docx)", 
            font=ctk.CTkFont(weight="bold"),
            command=self._select_prosit,
            fg_color="#E07A5F", hover_color="#D16043"
        )
        self.load_prosit_btn.pack(fill="x")

        # Scrollable Area for Revision
        self.scrollable_frame = ctk.CTkScrollableFrame(self, label_text="Révision des données")
        self.scrollable_frame.pack(fill="both", expand=True, padx=20, pady=10)

        for section_name in SECTION_KEYWORDS.keys():
            section_label = ctk.CTkLabel(self.scrollable_frame, text=section_name, font=ctk.CTkFont(weight="bold"))
            section_label.pack(anchor="w", padx=10, pady=(10, 2))
            
            text_area = ctk.CTkTextbox(self.scrollable_frame, height=100)
            text_area.pack(fill="x", padx=10, pady=(0, 5))
            self.text_widgets[section_name] = text_area

        # Generate Button
        self.generate_btn = ctk.CTkButton(
            self, text="📝 Générer le CER", 
            font=ctk.CTkFont(size=18, weight="bold"),
            height=50,
            command=self._generate_cer,
            fg_color="#2E8B57", hover_color="#1F5F3A"
        )
        self.generate_btn.pack(fill="x", padx=20, pady=20)

    def _save_filename_setting(self, *args):
        settings = load_settings()
        settings["default_filename"] = self.default_filename.get()
        save_settings(settings)

    def _select_template(self):
        f = filedialog.askopenfilename(title="Sélectionner le Template CER", filetypes=[("Word files", "*.docx")])
        if f:
            self.template_path.set(f)
            settings = load_settings()
            settings["template_path"] = f
            save_settings(settings)

    def _select_prosit(self):
        f = filedialog.askopenfilename(title="Sélectionner le Prosit Aller", filetypes=[("Word files", "*.docx")])
        if f:
            threading.Thread(target=self._parse_prosit, args=(f,), daemon=True).start()

    def _parse_prosit(self, file_path):
        try:
            doc = Document(file_path)
            # Store found sections and their priority
            found_sections = {s: {"content": [], "has_colon": False} for s in SECTION_KEYWORDS.keys()}
            current_section = None

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text: continue
                
                lower_text = text.lower().replace(" ", " ").strip()
                
                # 1. Check if this is an EXCLUDED section
                is_excl = False
                for ex_kw in EXCLUDED_KEYWORDS:
                    if (lower_text == ex_kw or lower_text.startswith(ex_kw + " ")) and len(text) < 60:
                        current_section = None
                        is_excl = True
                        break
                if is_excl: continue

                # 2. Check if this is a VALID section header
                found_header = False
                for section, kws in SECTION_KEYWORDS.items():
                    is_h, with_colon = is_header(text, kws)
                    if is_h and len(text) < 60:
                        # Priority logic: if we find a version with a colon, it overrides
                        if with_colon and not found_sections[section]["has_colon"]:
                            found_sections[section]["content"] = [] # Restart with better header
                            found_sections[section]["has_colon"] = True
                        
                        current_section = section
                        found_header = True
                        break
                
                if not found_header and current_section:
                    found_sections[current_section]["content"].append(text)

            # Prepare final content
            final_content = {s: d["content"] for s, d in found_sections.items()}
            self.after(0, lambda: self._update_revision_fields(final_content))
            
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Erreur", f"Erreur lors de la lecture du Prosit : {str(e)}"))

    def _update_revision_fields(self, content):
        for section, lines in content.items():
            if section in self.text_widgets:
                self.text_widgets[section].delete("1.0", "end")
                self.text_widgets[section].insert("1.0", "\n".join(lines))
        messagebox.showinfo("Succès", "Données extraites avec succès. Veuillez vérifier avant de générer.")

    def _generate_cer(self):
        t_path = self.template_path.get()
        if not t_path or not os.path.exists(t_path):
            messagebox.showwarning("Attention", "Veuillez d'abord sélectionner un template CER.")
            return

        default_name = self.default_filename.get() if self.default_filename.get() else "CER_Genere"
        save_path = filedialog.asksaveasfilename(
            title="Enregistrer le CER", 
            defaultextension=".docx",
            filetypes=[("Word files", "*.docx")],
            initialfile=f"{default_name}.docx"
        )
        
        if not save_path:
            return

        try:
            doc = Document(t_path)
            
            # Prepare replacements dictionary
            replacements = {}
            for section_name, widget in self.text_widgets.items():
                val = widget.get("1.0", "end-1c").strip()
                replacements[section_name] = val

            replace_content_by_titles(doc, replacements)
            doc.save(save_path)
            
            messagebox.showinfo("Succès", f"CER généré avec succès !\n\nSauvegardé ici : {save_path}")
            
            # Open the file
            os.startfile(save_path) if os.name == 'nt' else None

        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de la génération du CER : {str(e)}")
